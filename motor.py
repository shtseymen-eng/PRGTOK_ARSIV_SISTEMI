# PRGTOK motor.py - ISOPA isim düzeltmeli güncel sürüm
import os
import re
import shutil
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path

import pandas as pd
import pdfplumber

# OCR desteği (opsiyonel) — kurulum: brew install tesseract tesseract-lang
try:
    import pytesseract
    from PIL import Image as _PILImage
    _PYTESSERACT_OK = True
except ImportError:
    _PYTESSERACT_OK = False
# Windows Tesseract yolu
if _PYTESSERACT_OK:
    try:
        pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    except Exception:
        pass
# Word (.docx) icin metin okuma destegi (opsiyonel) -- kurulum: pip install python-docx
try:
    from docx import Document as _DocxDocument
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

OCR_TIMEOUT = 25              # OCR icin maksimum bekleme suresi (saniye) - takilan dosya tarama akisini kilitlemesin

ISLEM_KODU = "S"             # Program tarafından otomatik tarandı/sınıflandırıldı işareti
ELLE_DUZELT_KODU = "E"       # Elle düzeltildi işareti: isim değişmez (kod hariç), sadece doğru klasöre taşınır
                              # Kullanım: dosya adının SONUNA boşluk + E ekle
                              # Örn: "Belge T9 ANA E.pdf" → T9 klasörüne taşınır

# Eski (geriye dönük uyumluluk) kodlar – yeni taramada S / E koduna çevrilir
ESKI_ISLEM_KODU = "PRGTOK"
ESKI_ELLE_DUZELT_KODU = "PRGT"

BELGE_KLASORLERI = {
    "TANK_BASINC_RAPORU": "TANK_BASINC_RAPORU",
    "ISOPA": "ISOPA",
    # T9 alt grupları – hepsi TEK fiziksel klasörde birleşir: "T9"
    "T9_GECICI": "T9",
    "T9_MUAYENE": "T9",
    "T9_ANA": "T9",
    # Eski T9 klasörü – aynı birleşik T9 klasörüne gider
    "T9_MUAYENE_SERTIFIKASI": "T9",
    "TRAFIK_SIGORTASI": "TRAFIK_SIGORTASI",
    "TEHLIKELI_MADDE_SIGORTASI": "TEHLIKELI_MADDE_SIGORTASI",
    "FENNI_MUAYENE": "FENNI_MUAYENE",
    "SIZDIRMAZLIK": "SIZDIRMAZLIK",
    "YUKSEKTE_CALISABILIR_SAGLIK_RAPORU": "YUKSEKTE_CALISABILIR_SAGLIK_RAPORU",
    "SRC5": "SRC5",
    "YABANCI_PLAKA": "YABANCI_PLAKA",
    "DIGER_BELGELER": "DIGER_BELGELER",
    # NOT: OKUNAMAYANLAR, FARKLI_FORMAT_DOSYALAR ve ISLEM_RAPORLARI
    # PREGATE_ARSIV dışında, seçilen ana klasörün içinde tutulur.
}

# Dosya adında kullanılan kısa kategori etiketleri
TUR_DISPLAY = {
    "TANK_BASINC_RAPORU":              "TANK BASINC",
    "T9_GECICI":                       "T9 GECICI",
    "T9_MUAYENE":                      "T9 MUAYENE",
    "T9_ANA":                          "T9",
    "T9_MUAYENE_SERTIFIKASI":          "T9",       # eski uyumluluk
    "ISOPA":                           "ISOPA",
    "TRAFIK_SIGORTASI":                "TRAFIK SIGORTASI",
    "TEHLIKELI_MADDE_SIGORTASI":       "TEHLIKELI MADDE SIGORTASI",
    "FENNI_MUAYENE":                   "FENNI MUAYENE",
    "SIZDIRMAZLIK":                    "SIZDIRMAZLIK",
    "YUKSEKTE_CALISABILIR_SAGLIK_RAPORU": "YC SAGLIK",
    "SRC5":                            "SRC5",
    "YABANCI_PLAKA":                   "YABANCI PLAKA",
    "DIGER_BELGELER":                  "DIGER",
}
# PRGT dosyaları için ters arama: "TANK BASINC" → "TANK_BASINC_RAPORU"
TUR_DISPLAY_REVERSE = {v.replace(" ", "_"): k for k, v in TUR_DISPLAY.items()}


PDF_EXT = {".pdf"}
IMG_EXT = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp", ".gif"}
DOCX_EXT = {".docx"}

TURK_PLAKA_RE = re.compile(r"\b(0[1-9]|[1-7][0-9]|8[01])\s*[A-ZÇĞİÖŞÜ]{1,3}\s*\d{2,4}\b", re.I)
TANK_NO_RE = re.compile(r"\b[A-Z]{4}\s?\d{6}[-\s]?\d\b", re.I)
SASI_RE = re.compile(r"\b[A-HJ-NPR-Z0-9]{17}\b", re.I)
DATE_RE_LIST = [
    re.compile(r"\b(\d{2})[./-](\d{2})[./-](\d{4})\b"),
    re.compile(r"\b(\d{4})[./-](\d{2})[./-](\d{2})\b"),
]
EN_DATE_RE = re.compile(r"\b(\d{1,2})[-\s](JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)[-\s](\d{2,4})\b", re.I)
CAPACITY_RE = re.compile(r"\b(\d{2,3}[\.,]?\d{3}|\d{4,6})\s*(?:LT|LITRE|LITER|L)\b", re.I)
YABANCI_PLAKA_ADAY_RE = re.compile(r"\b[A-Z]{1,3}\s?\d{2,5}\s?[A-Z]{1,3}\b|\b[A-Z]{2}\s?\d{3,6}\b|\b\d{3,6}\s?[A-Z]{2,4}\b", re.I)


def temizle_ad(text, max_len=80):
    if not text:
        return "BILGIYOK"
    tr_map = str.maketrans("ÇĞİÖŞÜçğıöşü", "CGIOSUcgiosu")
    text = str(text).replace("\u00a0", " ").translate(tr_map).upper().strip()
    text = re.sub(r"[^A-Z0-9._-]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text[:max_len] if text else "BILGIYOK"


def anlamli_kimlik_mi(text):
    if not text:
        return False
    t = temizle_ad(text, 80)
    kotu = {"ON", "DELIVERED", "DELIVERED_ON", "CERTIFICATE", "DRIVER", "ISOPA", "BILGIYOK"}
    return t not in kotu and len(t) >= 5


def tarih_formatla(tarih):
    if not tarih:
        return ""
    for ix, rgx in enumerate(DATE_RE_LIST):
        m = rgx.search(tarih)
        if m:
            if ix == 0:
                gun, ay, yil = m.groups()
            else:
                yil, ay, gun = m.groups()
            return f"{gun}.{ay}.{yil}"
    return ""


def en_date_formatla(tarih):
    aylar = {
        "JAN": "01", "FEB": "02", "MAR": "03", "APR": "04", "MAY": "05", "JUN": "06",
        "JUL": "07", "AUG": "08", "SEP": "09", "OCT": "10", "NOV": "11", "DEC": "12",
    }
    m = EN_DATE_RE.search(tarih or "")
    if not m:
        return ""
    gun, ay, yil = m.groups()
    if len(yil) == 2:
        yil = "20" + yil
    return f"{gun.zfill(2)}.{aylar.get(ay.upper(), '00')}.{yil}"


def dnv_next_due_bul(metin):
    for pattern in [
        r"DATE\s+NEXT\s+INSPECTION\s+DUE\s*[:\-]?\s*(0?[1-9]|1[0-2])[/.-](\d{2})",
        r"NEXT\s+INSPECTION\s+DUE\s*[:\-]?\s*(0?[1-9]|1[0-2])[/.-](\d{2})",
    ]:
        m = re.search(pattern, metin.upper())
        if m:
            ay, yil = m.groups()
            return f"{ay.zfill(2)}.20{yil}"
    return ""


def tse_next_inspection_bul(metin):
    m = re.search(r"SONRAKI\s+MUAYENE.*?(\d{2})[/.-](20\d{2})", metin.upper(), re.S)
    if not m:
        m = re.search(r"NEXT\s+INSPECTION.*?(\d{2})[/.-](20\d{2})", metin.upper(), re.S)
    if m:
        ay, yil = m.groups()
        return f"{ay}.{yil}"
    return ""


def tum_tarihleri_bul(metin):
    bulunan = []
    for rgx in DATE_RE_LIST:
        for m in rgx.finditer(metin):
            bulunan.append(tarih_formatla(m.group(0)))
    for m in EN_DATE_RE.finditer(metin):
        bulunan.append(en_date_formatla(m.group(0)))
    return [x for x in bulunan if x]


def en_mantikli_tarih(metin, tur=""):
    if tur == "TANK_BASINC_RAPORU":
        x = dnv_next_due_bul(metin)
        if x:
            return x
    if tur == "T9_MUAYENE_SERTIFIKASI":
        x = tse_next_inspection_bul(metin)
        if x:
            return x

    patterns = [
        r"SON\s+GEÇERLİLİK\s+TARİHİ\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"SON\s+GECERLILIK\s+TARIHI\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"VALID\s+UNTIL\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"BİTİŞ\s+TARİHİ\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"BITIS\s+TARIHI\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"EXPIRATION\s+DATE\s*[:\-]?\s*(\d{4}[./-]\d{2}[./-]\d{2})",
        r"MUAYENE\s+GEÇERLİLİK\s+TARİHİ\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"MUAYENE\s+GECERLILIK\s+TARIHI\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"BİR\s+SONRAKİ\s+KONTROL\s+TARİHİ\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
        r"BIR\s+SONRAKI\s+KONTROL\s+TARIHI\s*[:\-]?\s*(\d{2}[./-]\d{2}[./-]\d{4})",
    ]
    for p in patterns:
        m = re.search(p, metin.upper(), re.S)
        if m:
            return tarih_formatla(m.group(1))

    tarihler = tum_tarihleri_bul(metin)
    if not tarihler:
        return ""
    try:
        normal = [x for x in tarihler if re.match(r"\d{2}\.\d{2}\.\d{4}", x)]
        sirali = sorted(set(normal), key=lambda x: datetime.strptime(x, "%d.%m.%Y"))
        return sirali[-1] if sirali else tarihler[-1]
    except Exception:
        return tarihler[-1]


def klasorleri_hazirla(ana_klasor):
    ana   = Path(ana_klasor)
    arsiv = ana / "PREGATE_ARSIV"
    arsiv.mkdir(exist_ok=True)
    for klasor in BELGE_KLASORLERI.values():
        (arsiv / klasor).mkdir(exist_ok=True)
    # PREGATE_ARSIV dışındaki klasörler
    for d in ("OKUNAMAYANLAR", "FARKLI_FORMAT_DOSYALAR", "ISLEM_RAPORLARI"):
        (ana / d).mkdir(exist_ok=True)
    return arsiv


def pdf_text_oku(dosya_yolu, max_sayfa=None):
    """PDF'ten metin çıkarır. max_sayfa verilirse sadece ilk N sayfa okunur
    (performans için) — hız kazandırır, büyük PDF'lerde tüm sayfaları okumaz."""
    metin = ""
    with pdfplumber.open(dosya_yolu) as pdf:
        sayfalar = pdf.pages if max_sayfa is None else pdf.pages[:max_sayfa]
        for sayfa in sayfalar:
            metin += "\n" + (sayfa.extract_text() or "")
    return metin.strip()


def resim_text_oku(dosya_yolu):
    """OCR ile resim dosyasından metin çıkar. Tesseract + Türkçe dil paketi gerektirir.
    Kurulum (Mac): brew install tesseract tesseract-lang

    NOT: pytesseract'a OCR_TIMEOUT (varsayılan 25sn) verilir. Bozuk/çok karmaşık
    bir görsel tesseract'ı sonsuza kadar bekletirse, süre aşımında işlem
    iptal edilir ve "" döndürülür — bu sayede tek bir kötü dosya tüm taramayı
    kilitlemez (program "donmaz")."""
    if not _PYTESSERACT_OK:
        return ""
    try:
        img = _PILImage.open(dosya_yolu)
        # Çok büyük görselleri küçült (OCR hızını ciddi şekilde artırır)
        MAX_BOYUT = 2200
        if max(img.size) > MAX_BOYUT:
            oran = MAX_BOYUT / max(img.size)
            yeni_boyut = (max(1, int(img.size[0] * oran)), max(1, int(img.size[1] * oran)))
            img = img.resize(yeni_boyut, _PILImage.LANCZOS)
        # Türkçe+İngilizce; süre aşımında veya hata durumunda sadece İngilizce dene.
        # Süre aşımı (RuntimeError) bir kez daha denenmez — çift bekleme olmasın.
        for lang in ("tur+eng", "eng"):
            try:
                return pytesseract.image_to_string(
                    img, lang=lang, config="--psm 6", timeout=OCR_TIMEOUT
                ).strip()
            except RuntimeError as e:
                if "timeout" in str(e).lower():
                    return ""
                continue
            except Exception:
                continue
        return ""
    except Exception:
        return ""


def docx_text_oku(dosya_yolu):
    """Word (.docx) belgesinden paragraf ve tablo hücrelerindeki metni çıkarır.
    Eski .doc (binary) formatı desteklenmez -- sadece .docx."""
    if not _DOCX_OK:
        return ""
    try:
        belge = _DocxDocument(dosya_yolu)
        parcalar = [p.text for p in belge.paragraphs]
        for tablo in belge.tables:
            for satir in tablo.rows:
                for hucre in satir.cells:
                    parcalar.append(hucre.text)
        return "\n".join(parcalar).strip()
    except Exception:
        return ""


def desteklenen_uzanti(uzanti):
    """Programın içerik okuyabildiği (PDF / resim+OCR / Word) bir uzantı mı?"""
    return uzanti in PDF_EXT or uzanti in IMG_EXT or uzanti in DOCX_EXT


def belge_metni_oku(dosya_yolu, uzanti, max_sayfa=None):
    """Uzantıya göre uygun okuyucuya yönlendirir. PDF için max_sayfa
    (performans sınırlaması) desteklenir; resim ve Word için yoksayılır.
    Desteklenmeyen uzantılarda "" döner."""
    if uzanti in PDF_EXT:
        return pdf_text_oku(dosya_yolu, max_sayfa=max_sayfa)
    if uzanti in IMG_EXT:
        return resim_text_oku(dosya_yolu)
    if uzanti in DOCX_EXT:
        return docx_text_oku(dosya_yolu)
    return ""


def plaka_bul(metin):
    m = TURK_PLAKA_RE.search(metin.upper())
    return temizle_ad(m.group(0).replace(" ", "")) if m else ""


def tank_no_bul(metin):
    m = TANK_NO_RE.search(metin.upper())
    if not m:
        return ""
    raw = m.group(0).replace(" ", "").upper()
    # XXXX999999-9 formatına normalize et (check digit'ten önce tire yoksa ekle)
    norm = re.sub(r"^([A-Z]{4})(\d{6})-?(\d)$", r"\1\2-\3", raw)
    return norm if norm else temizle_ad(raw)


def sasi_bul(metin):
    m = SASI_RE.search(metin.upper())
    return temizle_ad(m.group(0)) if m else ""


def kapasite_bul(metin):
    temiz = metin.upper().replace(".", "").replace(",", "")
    for pattern in [
        r"CAPACITY\s*\(L\)\s*[:\-]?\s*(\d{4,6})",
        r"KAPASITE\s*\(L\).*?TOPLAM\s*\(TOTAL\)\s*[:\-]?\s*(\d{4,6})",
        r"TOPLAM\s*\(TOTAL\)\s*[:\-]?\s*(\d{4,6})",
        r"TANK\s+HACMI\s*[:\-]?\s*(\d{4,6})\s*LT",
    ]:
        m = re.search(pattern, temiz, re.S)
        if m:
            return f"{m.group(1)}LT"
    m = CAPACITY_RE.search(temiz)
    if not m:
        return ""
    sayi = re.sub(r"\D", "", m.group(1))
    return f"{sayi}LT" if sayi else ""


def tank_kodu_bul(metin):
    for pattern in [
        r"TANK\s+KODU.*?(L[A-Z0-9]{3})",
        r"PORTABLE\s+TANK\s+INSTRUCTION.*?(L[A-Z0-9]{3})",
    ]:
        m = re.search(pattern, metin.upper(), re.S)
        if m:
            return temizle_ad(m.group(1))
    m = re.search(r"\bL[A-Z0-9]{3}\b", metin.upper())
    return temizle_ad(m.group(0)) if m else ""


def un_numaralari_bul(metin, max_adet=10):
    """Metindeki UN numaralarını çıkarır. Dönüş: '1090-2055-1302' formatında string."""
    eslesler = re.findall(r"\bUN\s*N?O?\.?\s*(\d{4})\b", metin.upper())
    gorulen = set()
    benzersiz = []
    for un in eslesler:
        if un not in gorulen:
            gorulen.add(un)
            benzersiz.append(un)
        if len(benzersiz) >= max_adet:
            break
    return "-".join(benzersiz)


def yabanci_plaka_bul(metin):
    temiz = TURK_PLAKA_RE.sub(" ", metin.upper())
    yasak = ("POL", "TANK", "ADR", "PDF", "DNV", "CSC", "RID", "IMDG", "DATE", "TEST",
             "NEXT", "OWNER", "SHELL", "GROSS", "PRESSURE", "CAPACITY", "INSPECTION",
             "SERIAL", "TSE", "ISOPA", "DRIVER", "CERTIFICATE")
    adaylar = []
    for m in YABANCI_PLAKA_ADAY_RE.finditer(temiz):
        aday = temizle_ad(m.group(0).replace(" ", ""))
        if 5 <= len(aday) <= 12 and not aday.startswith(yasak):
            adaylar.append(aday)
    return adaylar[0] if adaylar else ""


def surucu_adi_bul(metin, dosya_adi=""):
    metin_norm = (metin or "").replace("\u00a0", " ")
    buyuk = metin_norm.upper()

    # ISOPA özel okuma: isim, "certificate to" ile "under the supervision" arasında kalır.
    m = re.search(
        r"ISOPA\s+DELIVERED\s+THIS\s+CERTIFICATE\s+TO\s+(.+?)(?:UNDER\s+THE\s+SUPERVISION|DELIVERED\s+ON|EXPIRATION\s+DATE|CERTIFICATE\s+UNIQUE\s+NUMBER|$)",
        buyuk,
        re.S
    )
    if m:
        isim = m.group(1)
        isim = re.sub(r"[^A-ZÇĞİÖŞÜ\s]", " ", isim)
        isim = re.sub(r"\s+", " ", isim).strip()
        if anlamli_kimlik_mi(isim):
            return temizle_ad(isim)

    # Alternatif ISOPA: TO sonrası tek satır isim yakala.
    m = re.search(r"TO\s+([A-ZÇĞİÖŞÜ]{2,}\s+[A-ZÇĞİÖŞÜ]{2,}(?:\s+[A-ZÇĞİÖŞÜ]{2,})?)", buyuk)
    if m:
        isim = m.group(1).strip()
        if anlamli_kimlik_mi(isim):
            return temizle_ad(isim)

    # Sağlık raporu / genel sürücü adı
    m = re.search(r"(?:ADI\s*SOYADI|AD SOYAD|NAME\s*SURNAME)\s*[:\-]?\s*([A-ZÇĞİÖŞÜ ]{5,60})", buyuk)
    if m:
        isim = m.group(1).strip()
        if anlamli_kimlik_mi(isim):
            return temizle_ad(isim)

    # Dosya adından fallback: Nihat YILMAZ- ISOPA.pdf gibi
    stem = Path(dosya_adi).stem
    stem_clean = re.sub(r"[-_]", " ", stem).upper()
    if any(x in stem_clean for x in ["ISOPA", "YÇ", "YC", "YUKSEKTE"]):
        stem_clean = re.sub(r"\b(ISOPA|YÇ|YC|YUKSEKTE|CALISABILIR|ÇALISABILIR|ÇALIŞABİLİR|RAPORU)\b", " ", stem_clean)
        stem_clean = re.sub(r"\d{2,4}[./-]\d{1,2}[./-]\d{1,4}", " ", stem_clean)
        stem_clean = re.sub(r"\s+", " ", stem_clean).strip()
        if anlamli_kimlik_mi(stem_clean):
            return temizle_ad(stem_clean)

    return ""


def bilgi_dosya_adindan(dosya_adi):
    stem = Path(dosya_adi).stem
    return {
        "plaka": plaka_bul(stem),
        "tank_no": tank_no_bul(stem),
        "sasi": sasi_bul(stem),
        "yabanci_plaka": yabanci_plaka_bul(stem),
        "tarih": en_mantikli_tarih(stem),
        "kapasite": kapasite_bul(stem),
        "tank_kodu": tank_kodu_bul(stem),
        "surucu": surucu_adi_bul("", dosya_adi),
    }


def belge_turu_bul(metin, dosya_adi=""):
    m = (metin + "\n" + dosya_adi).upper()
    ad = dosya_adi.upper()

    if "YÇ" in ad or "YC" in ad or "YUKSEKTE" in m or "YÜKSEKTE" in m:
        return "YUKSEKTE_CALISABILIR_SAGLIK_RAPORU"

    if ("TEHLIKELI MADDE" in m or "TEHLİKELİ MADDE" in m or "TEHLIKELI ATIK" in m or "TEHLİKELİ ATIK" in m) and "ZORUNLU MALI SORUMLULUK" in m:
        return "TEHLIKELI_MADDE_SIGORTASI"

    # ── T9 GEÇİCİ: 23 iş günü geçerli geçici belgeler ──────────────────────────
    if ("GEÇİCİ" in m or "GECICI" in m or
        "TEMPORARY" in m or "PROVISIONAL" in m or "INTERIM" in m or
        "23 İŞ GÜNÜ" in m or "23 IS GUNU" in m or "23 WORKING DAY" in m) and (
        "T9" in m or "ADR" in m or "UYGUNLUK" in m or "ONAY" in m or "APPROVAL" in m):
        return "T9_GECICI"

    # ── T9 MUAYENE: Muayene sertifikası – şasi no + litre/kapasite ───────────
    if ("MUAYENE SERTIFIKASI" in m or "MUAYENE SERTİFİKASI" in m) and (
        sasi_bul(m) or "ŞASİ" in m or "SASI NO" in m or "VIN" in m
    ) and (
        kapasite_bul(m) or "KAPASİTE" in m or "KAPASITE" in m or "LİTRE" in m or "LITRE" in m
    ):
        return "T9_MUAYENE"

    # ── T9 ANA: Asıl ADR onay belgesi – tank kodu, plaka, UN numaraları ──────
    if ("BELİRLİ TEHLİKELİ MADDELER TAŞIYAN ARAÇLAR İÇİN ONAY SERTİFİKASI" in m or
        "BELIRLI TEHLIKELI MADDELER TASIYAN ARACLAR ICIN ONAY SERTIFIKASI" in m or
        "CERTIFICATE OF APPROVAL FOR VEHICLES CARRYING CERTAIN DANGEROUS GOODS" in m or
        "TAŞIT UYGUNLUK" in m or "TASIT UYGUNLUK" in m or "ADR UYGUNLUK" in m or
        "PORTABLE TANK INSTRUCTION" in m or
        "TANK BİLGİLERİ" in m or "TANK BILGILERI" in m or
        "TANK KODU" in m or "TASIMA BIRIM TIPI" in m or "TAŞIMA BIRIM TIPI" in m or
        bool(re.search(r"\bUN\s*\d{4}\b", m)) or
        ("INSPECTION CERTIFICATE" in m and (
            "TANK" in m or "CONTAINER" in m or "ADR" in m
        )) or
        ("T9" in m and ("ADR" in m or "ONAY SERTIFIKA" in m or "ONAY SERTİFİKA" in m or
                        "UYGUNLUK BELGE" in m or "APPROVAL" in m))):
        return "T9_ANA"

    # ── SRC-5: Tehlikeli madde taşımacılığı sürücü sertifikası ───────────────
    if ("SRC-5" in m or "SRC 5" in m or "SRC5" in m or
        ("SRC" in m and ("TEHLIKELI MADDE TASIMACILIGI" in m or "TEHLİKELİ MADDE TAŞIMACILIĞI" in m or
                         "DANGEROUS GOODS DRIVER" in m or "ADR SURUCU" in m or "ADR SÜRÜCÜ" in m))):
        return "SRC5"

    # ── TANK BASINÇ RAPORU: Periyodik muayene / basınç test raporları ─────────
    tank_ifadeleri = [
        # DNV GL / DNV
        "TANK CONTAINER PERIODIC INSPECTION REPORT", "PERIODIC INSPECTION REPORT",
        "DATE NEXT INSPECTION DUE", "NEXT INSPECTION DUE",
        "DNV SILVER", "SILVER/CIMS", "DNV GL",
        # Genel İngilizce
        "OWNER'S SERIAL NUMBER", "OWNERS SERIAL NUMBER", "OWNER S SERIAL NUMBER",
        "OWNER'S SERIAL NO", "OWNERS SERIAL NO",
        "INITIAL HYDRO TEST", "LAST HYDRO TEST", "HYDROSTATIC TEST",
        "THIS INSPECTION", "INSPECTION DATES", "PERIODIC TEST",
        "NEXT PERIODIC TEST", "NEXT TEST DATE", "RETEST DATE",
        "MAX GROSS WEIGHT", "TARE WEIGHT",
        "CAPACITY (L)", "CAPACITY(L)",
        "TEST PRESSURE", "DESIGN PRESSURE", "WORKING PRESSURE",
        "M.A.W.P", "MAWP",
        "ISO TYPE", "PRESSURE RELIEF VALVES",
        "SHELL THICKNESS", "MINIMUM DESIGN METAL TEMPERATURE", "MDMT",
        # Muayene firmaları
        "BUREAU VERITAS", "INTERTEK", "SGS INSPECTION", "COTAC",
        "RINA INSPECTION", "ABS INSPECTION", "TUV INSPECTION",
        # Türkçe
        "BASINÇ TEST RAPORU", "BASINC TEST RAPORU",
        "PERİYODİK MUAYENE RAPORU", "PERIYODIK MUAYENE RAPORU",
    ]
    if any(x in m for x in tank_ifadeleri):
        return "TANK_BASINC_RAPORU"

    if tank_no_bul(m) and ("PRESSURE" in m or "BASINC" in m or "BASINÇ" in m or
                            "HYDRO" in m or "TANK CONTAINER" in m or "PERIODIC" in m):
        return "TANK_BASINC_RAPORU"

    if "ISOPA" in m or "TDI" in m or "MDI" in m:
        return "ISOPA"

    if "TRAFIK SIGORT" in m or "TRAFİK SİGORT" in m or "KARAYOLLARI MOTORLU" in m:
        return "TRAFIK_SIGORTASI"

    if "FENNI MUAYENE" in m or "FENNİ MUAYENE" in m or "ARAÇ MUAYENE RAPORU" in m or "ARAC MUAYENE RAPORU" in m or "VEHICLE INSPECTION REPORT" in m:
        return "FENNI_MUAYENE"

    if ("SIZDIRMAZLIK RAPORU" in m or "SIZDIRMAZLIK TEST RAPORU" in m or "KARA TANKERI SIZDIRMAZLIK TEST RAPORU" in m or
        "LEAKPROOFNESS TEST REPORT" in m or "LEAK PROOFNESS TEST REPORT" in m or "LEAKPROOFNESS CERTIFICATE" in m):
        return "SIZDIRMAZLIK"

    if yabanci_plaka_bul(m):
        return "YABANCI_PLAKA"

    return "DIGER_BELGELER"



def _son_token(dosya_adi):
    """Dosya adının (uzantısız) son boşlukla ayrılmış parçasını büyük harfle döndürür.
    İşlem kodu tespiti SADECE bu son parçaya bakılarak yapılır; bu sayede
    tek harfli kodlar (S / E) dosya adı içinde başka yerlerde geçse bile
    yanlış pozitif (false positive) oluşmaz."""
    stem = Path(dosya_adi).stem.strip()
    if not stem:
        return ""
    parcalar = stem.replace("-", " ").split(" ")
    return parcalar[-1].upper().strip() if parcalar else ""


def _islenmis_mi(dosya_adi):
    """Dosya program tarafından otomatik tarandı mı? (yeni kod 'S' veya eski 'PRGTOK')"""
    son = _son_token(dosya_adi)
    return son in (ISLEM_KODU, ESKI_ISLEM_KODU)


def _is_prgt_dosya(dosya_adi):
    """Dosya elle düzeltme kodu ile mi işaretli? (yeni kod 'E' veya eski 'PRGT')
    Kullanım: elle düzeltilmiş belgeler için — sadece doğru klasöre taşınır,
    kodu yeni taramada 'E' olarak güncellenir."""
    son = _son_token(dosya_adi)
    return son in (ELLE_DUZELT_KODU, ESKI_ELLE_DUZELT_KODU)


def _kod_degistir(dosya_adi, yeni_kod):
    """Dosya adının sonundaki işlem kodunu (S/PRGTOK/E/PRGT) yeni_kod ile değiştirir.
    Son parça bilinen bir kod değilse, yeni_kod sona eklenir."""
    p = Path(dosya_adi)
    stem = p.stem.strip()
    bilinen_kodlar = (ISLEM_KODU, ELLE_DUZELT_KODU, ESKI_ISLEM_KODU, ESKI_ELLE_DUZELT_KODU)
    parcalar = stem.split(" ")
    if parcalar and parcalar[-1].upper() in bilinen_kodlar:
        parcalar[-1] = yeni_kod
    else:
        parcalar.append(yeni_kod)
    return " ".join(x for x in parcalar if x) + p.suffix


def belge_turu_dosya_adindan(dosya_adi):
    # Boşluklu ve alt tireli isimlerin ikisini de tanımak için normalize et
    ad = dosya_adi.upper().replace(" ", "_")

    # Eski T9_MUAYENE_SERTIFIKASI dosyalarını T9_ANA'ya yönlendir
    if "T9_MUAYENE_SERTIFIKASI" in ad or "T9_SERTIFIKASI" in ad:
        return "T9_ANA"

    for tur in BELGE_KLASORLERI.keys():
        if tur not in ("ISLEM_RAPORLARI", "T9_MUAYENE_SERTIFIKASI") and tur in ad:
            return tur

    # PRGT dosyaları için TUR_DISPLAY kısa adlarıyla ters arama
    # Örn: "TANK_BASINC" → "TANK_BASINC_RAPORU"; "T9" → "T9_ANA"
    for display_norm, tur in TUR_DISPLAY_REVERSE.items():
        if display_norm and len(display_norm) >= 2 and display_norm in ad:
            return tur

    if _islenmis_mi(dosya_adi) and YABANCI_PLAKA_ADAY_RE.search(ad) and not TURK_PLAKA_RE.search(ad):
        return "YABANCI_PLAKA"
    return ""


def benzersiz_yol(hedef_yol):
    if not hedef_yol.exists():
        return hedef_yol
    stem, suffix, parent = hedef_yol.stem, hedef_yol.suffix, hedef_yol.parent
    i = 1
    while True:
        aday = parent / f"{stem}_{i}{suffix}"
        if not aday.exists():
            return aday
        i += 1


_TASIMA_KILIDI = threading.Lock()


def dosya_tasi(kaynak, hedef_klasor, yeni_ad=None):
    """Dosyayı hedef klasöre taşır. Paralel tarama sırasında aynı klasöre
    aynı anda yazma/isim çakışması olmaması için kilit kullanılır.

    NOT: Hedef yol zaten kaynağın kendisiyse (dosya zaten doğru yerde ve
    doğru isimde) hiçbir şey yapılmaz — benzersiz_yol() bu durumda dosyanın
    kendisini "doluyor" sanıp gereksiz "_1" son eki eklemesin diye bu kontrol
    benzersiz_yol'dan ÖNCE yapılır."""
    with _TASIMA_KILIDI:
        hedef_klasor.mkdir(parents=True, exist_ok=True)
        hedef_aday = hedef_klasor / (yeni_ad if yeni_ad else kaynak.name)
        if kaynak.resolve() == hedef_aday.resolve():
            return hedef_aday
        hedef = benzersiz_yol(hedef_aday)
        shutil.move(str(kaynak), str(hedef))
        return hedef


def _fp(text, max_len=0):
    """Dosya adı parçası: sadece A-Z 0-9 nokta izinli, diğerleri boşluk.
    Tank kodu ve UN numara listesi gibi tire içeren değerler bu fonksiyona GEÇİRİLMEZ."""
    if not text:
        return ""
    tr_map = str.maketrans("ÇĞİÖŞÜçğıöşü", "CGIOSUcgiosu")
    t = str(text).replace("\u00a0", " ").translate(tr_map).upper().strip()
    t = re.sub(r"[^A-Z0-9.]+", " ", t)
    t = re.sub(r" +", " ", t).strip()
    return t[:max_len] if max_len else t


def yeni_dosya_adi_olustur(tur, bilgi, eski_ad, uzanti):
    tank_no      = bilgi.get("tank_no") or ""
    tarih        = bilgi.get("tarih") or ""
    kapasite     = bilgi.get("kapasite") or ""
    tank_kodu    = bilgi.get("tank_kodu") or ""
    un_numaralar = bilgi.get("un_numaralar") or ""
    tur_ad       = TUR_DISPLAY.get(tur, tur.replace("_", " "))

    if tur == "ISOPA":
        kimlik = bilgi.get("surucu") or temizle_ad(Path(eski_ad).stem, 40)
    else:
        kimlik = (
            tank_no or bilgi.get("plaka") or bilgi.get("sasi") or
            bilgi.get("yabanci_plaka") or bilgi.get("surucu") or temizle_ad(Path(eski_ad).stem, 40)
        )
    if not anlamli_kimlik_mi(kimlik):
        kimlik = temizle_ad(Path(eski_ad).stem, 40)

    # Tank kodu (CAIU123456-7) → tire korunur; diğer tüm kimlikler → _fp ile tire kaldırılır
    kimlik_part = tank_no if (tank_no and kimlik == tank_no) else _fp(kimlik, 60)

    # ── YABANCI PLAKA ──────────────────────────────────────────────────────────
    if tur == "YABANCI_PLAKA":
        parcalar = [_fp(bilgi.get("yabanci_plaka") or kimlik_part)]
        if tarih:
            parcalar.append(tarih)
        parcalar.append(ISLEM_KODU)
        return " ".join(p for p in parcalar if p)[:160] + uzanti.lower()

    # ── TANK BASINÇ: tank_no  tarih  tank_kodu  kapasite  TANK BASINC  PRGTOK ─
    if tur == "TANK_BASINC_RAPORU":
        parcalar = [kimlik_part]
        if tarih:     parcalar.append(tarih)
        if tank_kodu: parcalar.append(_fp(tank_kodu))
        if kapasite:  parcalar.append(_fp(kapasite))
        parcalar += [tur_ad, ISLEM_KODU]
        return " ".join(p for p in parcalar if p)[:180] + uzanti.lower()

    # ── T9 ANA: kimlik  tarih  kapasite  tank_kodu  T9  UN-numaralar  PRGTOK ──
    if tur in ("T9_ANA", "T9_MUAYENE_SERTIFIKASI"):
        parcalar = [kimlik_part]
        if tarih:        parcalar.append(tarih)
        if kapasite:     parcalar.append(_fp(kapasite))
        if tank_kodu:    parcalar.append(_fp(tank_kodu))
        parcalar.append(tur_ad)
        if un_numaralar: parcalar.append(un_numaralar)   # 1090-2055-1302 – tireler korunur
        parcalar.append(ISLEM_KODU)
        return " ".join(p for p in parcalar if p)[:180] + uzanti.lower()

    # ── DİĞER TÜM TÜRLER: kimlik  tarih  [kapasite]  tur_ad  PRGTOK ───────────
    parcalar = [kimlik_part]
    if tarih:   parcalar.append(tarih)
    if kapasite and tur in ("T9_MUAYENE", "SIZDIRMAZLIK"):
        parcalar.append(_fp(kapasite))
    parcalar += [tur_ad, ISLEM_KODU]
    return " ".join(p for p in parcalar if p)[:180] + uzanti.lower()


def bilgi_cek(metin, dosya_adi=""):
    tur = belge_turu_bul(metin, dosya_adi)
    dosya_bilgisi = bilgi_dosya_adindan(dosya_adi)
    return {
        "plaka": plaka_bul(metin) or dosya_bilgisi.get("plaka", ""),
        "tank_no": tank_no_bul(metin) or dosya_bilgisi.get("tank_no", ""),
        "sasi": sasi_bul(metin) or dosya_bilgisi.get("sasi", ""),
        "yabanci_plaka": yabanci_plaka_bul(metin) or dosya_bilgisi.get("yabanci_plaka", ""),
        "tarih": en_mantikli_tarih(metin, tur) or dosya_bilgisi.get("tarih", ""),
        "kapasite": kapasite_bul(metin) or dosya_bilgisi.get("kapasite", ""),
        "tank_kodu": tank_kodu_bul(metin) or dosya_bilgisi.get("tank_kodu", ""),
        "surucu": surucu_adi_bul(metin, dosya_adi) or dosya_bilgisi.get("surucu", ""),
        "un_numaralar": un_numaralari_bul(metin),
    }


def hatali_isopa_prgtok_mu(dosya_adi):
    ad = dosya_adi.upper().replace(" ", "_")
    return _islenmis_mi(dosya_adi) and "_ISOPA_" in ad and (
        ad.startswith("ON_") or ad.startswith("DELIVERED_") or ad.startswith("BILGIYOK_") or ad.startswith("CERTIFICATE_")
    )


def dosya_isle(dosya_yolu, ana_klasor):
    kaynak = Path(dosya_yolu)
    arsiv = klasorleri_hazirla(ana_klasor)
    eski_ad, uzanti = kaynak.name, kaynak.suffix.lower()

    sonuc = {
        "Eski Dosya Adı": eski_ad, "Yeni Dosya Adı": "", "Belge Türü": "", "Plaka": "", "Tank No": "",
        "Şasi No": "", "Yabancı Plaka": "", "Geçerlilik Tarihi": "", "Kapasite": "", "Tank Kodu": "",
        "Sürücü": "", "Bulunduğu Klasör": str(kaynak.parent), "Yeni Klasör": "", "İşlem Durumu": "",
        "Hata": "", "İşlem Tarihi": datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
    }

    try:
        if kaynak.name.startswith("~$"):
            sonuc["İşlem Durumu"] = "GECICI_DOSYA_ATLANDI"
            return sonuc

        if "ISLEM_RAPORLARI" in str(kaynak.parent):
            sonuc["İşlem Durumu"] = "RAPOR_ATLANDI"
            return sonuc

        # ── E (elle düzeltilmiş) / eski PRGT: isim değişmez (kod hariç), sadece doğru klasöre at ──
        if _is_prgt_dosya(kaynak.name):
            # İçeriği oku → türü belirle (dosya adı fallback)
            tur = ""
            if desteklenen_uzanti(uzanti):
                try:
                    _m = belge_metni_oku(str(kaynak), uzanti)
                    tur = belge_turu_bul(_m, kaynak.name)
                except Exception:
                    pass
            if not tur or tur == "DIGER_BELGELER":
                tur = belge_turu_dosya_adindan(kaynak.name) or "DIGER_BELGELER"
            hedef_klasor = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
            sonuc["Belge Türü"] = tur
            # Eski " PRGT" kodunu yeni "E" koduna çevir
            yeni_ad = kaynak.name if _son_token(kaynak.name) == ELLE_DUZELT_KODU else _kod_degistir(kaynak.name, ELLE_DUZELT_KODU)
            if kaynak.parent.resolve() != hedef_klasor.resolve() or yeni_ad != kaynak.name:
                yeni_yol = dosya_tasi(kaynak, hedef_klasor, yeni_ad)
                sonuc["Yeni Dosya Adı"] = yeni_yol.name
                sonuc["Yeni Klasör"] = str(yeni_yol.parent)
                sonuc["İşlem Durumu"] = "ELLEDUZ_TASINDI"
            else:
                sonuc["Yeni Dosya Adı"] = kaynak.name
                sonuc["Yeni Klasör"] = str(kaynak.parent)
                sonuc["İşlem Durumu"] = "ELLEDUZ_DOGRU_YERDE"
            return sonuc

        # ── S (otomatik tarandı) / eski PRGTOK: Daha önce işlenmiş dosyalar ─────
        if _islenmis_mi(kaynak.name) and not hatali_isopa_prgtok_mu(kaynak.name):
            tur = belge_turu_dosya_adindan(kaynak.name)

            # OKUNAMAYANLAR → yeniden okuma dene
            if tur == "OKUNAMAYANLAR" and desteklenen_uzanti(uzanti):
                try:
                    _m = belge_metni_oku(str(kaynak), uzanti)
                except Exception:
                    _m = ""
                if _m:
                    _tur  = belge_turu_bul(_m, kaynak.name)
                    _bilgi = bilgi_cek(_m, kaynak.name)
                    _ad   = yeni_dosya_adi_olustur(_tur, _bilgi, kaynak.name, uzanti)
                    _yol  = dosya_tasi(kaynak, arsiv / BELGE_KLASORLERI.get(_tur, "DIGER_BELGELER"), _ad)
                    sonuc.update({
                        "Yeni Dosya Adı": _yol.name,
                        "Belge Türü": _tur,
                        "Plaka": _bilgi.get("plaka", ""),
                        "Tank No": _bilgi.get("tank_no", ""),
                        "Şasi No": _bilgi.get("sasi", ""),
                        "Yabancı Plaka": _bilgi.get("yabanci_plaka", ""),
                        "Geçerlilik Tarihi": _bilgi.get("tarih", ""),
                        "Kapasite": _bilgi.get("kapasite", ""),
                        "Tank Kodu": _bilgi.get("tank_kodu", ""),
                        "Sürücü": _bilgi.get("surucu", ""),
                        "Yeni Klasör": str(_yol.parent),
                        "İşlem Durumu": "OKUNAMAYAN_OKUNDU_TASINDI",
                    })
                    return sonuc
                # Hâlâ okunamıyor → OKUNAMAYANLAR klasöründe bırak
                hedef_klasor = arsiv.parent / "OKUNAMAYANLAR"
                if kaynak.parent.resolve() != hedef_klasor.resolve():
                    yeni_yol = dosya_tasi(kaynak, hedef_klasor)
                    sonuc.update({"Yeni Dosya Adı": yeni_yol.name,
                                  "Yeni Klasör": str(yeni_yol.parent),
                                  "İşlem Durumu": "HALA_OKUNAMADI_TASINDI"})
                else:
                    sonuc.update({"Yeni Dosya Adı": kaynak.name,
                                  "Yeni Klasör": str(kaynak.parent),
                                  "İşlem Durumu": "HALA_OKUNAMADI"})
                sonuc["Belge Türü"] = "OKUNAMAYANLAR"
                return sonuc

            if not tur:
                sonuc["İşlem Durumu"] = "ATLANDI_PRGTOK_TUR_BULUNAMADI"
                return sonuc

            hedef_klasor = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
            sonuc["Belge Türü"] = tur
            # Eski "PRGTOK" kodunu yeni "S" koduna çevir
            yeni_ad = kaynak.name if _son_token(kaynak.name) == ISLEM_KODU else _kod_degistir(kaynak.name, ISLEM_KODU)
            if kaynak.parent.resolve() != hedef_klasor.resolve() or yeni_ad != kaynak.name:
                yeni_yol = dosya_tasi(kaynak, hedef_klasor, yeni_ad)
                sonuc["Yeni Dosya Adı"] = yeni_yol.name
                sonuc["Yeni Klasör"] = str(yeni_yol.parent)
                sonuc["İşlem Durumu"] = "YER_DUZELTILDI_PRGTOK"
            else:
                sonuc["Yeni Dosya Adı"] = kaynak.name
                sonuc["Yeni Klasör"] = str(kaynak.parent)
                sonuc["İşlem Durumu"] = "ATLANDI_PRGTOK"
            return sonuc

        # ── Desteklenmeyen format ────────────────────────────────────────────────
        if not desteklenen_uzanti(uzanti):
            yeni_yol = dosya_tasi(kaynak, arsiv.parent / "FARKLI_FORMAT_DOSYALAR")
            sonuc.update({
                "Belge Türü": "FARKLI_FORMAT_DOSYALAR",
                "Yeni Dosya Adı": yeni_yol.name,
                "Yeni Klasör": str(yeni_yol.parent),
                "İşlem Durumu": "FARKLI_FORMAT_TASINDI"
            })
            return sonuc

        # ── PDF / Resim oku ──────────────────────────────────────────────────────
        try:
            if uzanti in IMG_EXT:
                metin = resim_text_oku(str(kaynak))
                if not metin and not _PYTESSERACT_OK:
                    yeni_yol = dosya_tasi(kaynak, arsiv.parent / "FARKLI_FORMAT_DOSYALAR")
                    sonuc.update({
                        "Belge Türü": "FARKLI_FORMAT_DOSYALAR",
                        "Yeni Dosya Adı": yeni_yol.name,
                        "Yeni Klasör": str(yeni_yol.parent),
                        "İşlem Durumu": "TESSERACT_KURULU_DEGIL",
                        "Hata": "pip install pytesseract + winget install UB-Mannheim.TesseractOCR",
                    })
                    return sonuc
            elif uzanti in DOCX_EXT:
                metin = docx_text_oku(str(kaynak))
                if not metin and not _DOCX_OK:
                    yeni_yol = dosya_tasi(kaynak, arsiv.parent / "FARKLI_FORMAT_DOSYALAR")
                    sonuc.update({
                        "Belge Türü": "FARKLI_FORMAT_DOSYALAR",
                        "Yeni Dosya Adı": yeni_yol.name,
                        "Yeni Klasör": str(yeni_yol.parent),
                        "İşlem Durumu": "PYTHON_DOCX_KURULU_DEGIL",
                        "Hata": "pip install python-docx",
                    })
                    return sonuc
            else:
                # Performans: önce ilk birkaç sayfa okunur; yetersizse tüm PDF okunur
                metin = pdf_text_oku(str(kaynak), max_sayfa=6)
        except Exception as e:
            metin = ""
            sonuc["Hata"] = f"OKUMA_HATASI: {e}"

        tur = belge_turu_bul(metin, eski_ad)

        # İlk sayfalarda yeterli bilgi bulunamadıysa PDF'in tamamını oku
        if uzanti in PDF_EXT and tur == "DIGER_BELGELER" and not yabanci_plaka_bul(metin):
            try:
                metin_tam = pdf_text_oku(str(kaynak))
                if len(metin_tam) > len(metin):
                    metin = metin_tam
                    tur = belge_turu_bul(metin, eski_ad)
            except Exception:
                pass

        bilgi = bilgi_cek(metin, eski_ad)

        # Metin okunamadıysa → orijinal adıyla OKUNAMAYANLAR klasörüne at (prefix yok)
        if not metin and tur == "DIGER_BELGELER":
            yeni_yol = dosya_tasi(kaynak, arsiv.parent / "OKUNAMAYANLAR")
            sonuc.update({
                "Belge Türü": "OKUNAMAYANLAR",
                "Yeni Dosya Adı": yeni_yol.name,
                "Yeni Klasör": str(yeni_yol.parent),
                "İşlem Durumu": "OKUNAMADI_TASINDI"
            })
            return sonuc

        yeni_ad = yeni_dosya_adi_olustur(tur, bilgi, eski_ad, uzanti)
        yeni_yol = dosya_tasi(kaynak, arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER"), yeni_ad)

        sonuc.update({
            "Yeni Dosya Adı": yeni_yol.name,
            "Belge Türü": tur,
            "Plaka": bilgi.get("plaka", ""),
            "Tank No": bilgi.get("tank_no", ""),
            "Şasi No": bilgi.get("sasi", ""),
            "Yabancı Plaka": bilgi.get("yabanci_plaka", ""),
            "Geçerlilik Tarihi": bilgi.get("tarih", ""),
            "Kapasite": bilgi.get("kapasite", ""),
            "Tank Kodu": bilgi.get("tank_kodu", ""),
            "Sürücü": bilgi.get("surucu", ""),
            "Yeni Klasör": str(yeni_yol.parent),
            "İşlem Durumu": "OKUNDU_TASINDI",
        })
        return sonuc

    except Exception as e:
        sonuc["İşlem Durumu"] = "HATA"
        sonuc["Hata"] = str(e)
        return sonuc


def _eski_konumlari_tasi(ana_klasor):
    """Eski PREGATE_ARSIV/OKUNAMAYANLAR ve PREGATE_ARSIV/FARKLI_FORMAT_DOSYALAR
    içindeki dosyaları yeni dış konumlarına taşır (tek seferlik geçiş)."""
    ana   = Path(ana_klasor)
    arsiv = ana / "PREGATE_ARSIV"
    gecis = {
        "OKUNAMAYANLAR":       ana / "OKUNAMAYANLAR",
        "FARKLI_FORMAT_DOSYALAR": ana / "FARKLI_FORMAT_DOSYALAR",
    }
    for klasor_adi, yeni_konum in gecis.items():
        eski_konum = arsiv / klasor_adi
        if not eski_konum.exists():
            continue
        yeni_konum.mkdir(exist_ok=True)
        for dosya in list(eski_konum.iterdir()):
            if dosya.is_file():
                dosya_tasi(dosya, yeni_konum, dosya.name)
        # Eski klasör boşaldıysa sil
        try:
            eski_konum.rmdir()
        except OSError:
            pass  # İçinde hâlâ dosya varsa bırak


def _eski_t9_klasorlerini_temizle(ana_klasor):
    """T9_GECICI / T9_MUAYENE / T9_ANA / T9_MUAYENE_SERTIFIKASI artık 'T9'
    klasörüyle birleştiği için, içi boşaldıysa eski klasörleri siler."""
    arsiv = Path(ana_klasor) / "PREGATE_ARSIV"
    for eski_ad in ("T9_GECICI", "T9_MUAYENE", "T9_ANA", "T9_MUAYENE_SERTIFIKASI"):
        eski_klasor = arsiv / eski_ad
        if eski_klasor.exists() and eski_klasor.is_dir():
            try:
                if not any(eski_klasor.iterdir()):
                    eski_klasor.rmdir()
            except OSError:
                pass


def klasor_tara(ana_klasor, log_callback=None):
    arsiv = klasorleri_hazirla(ana_klasor)
    sonuclar = []
    ana = Path(ana_klasor)

    # Eski PREGATE_ARSIV içindeki OKUNAMAYANLAR/FARKLI_FORMAT dosyalarını dışarı taşı
    _eski_konumlari_tasi(ana_klasor)
    if log_callback:
        log_callback("Eski klasor yapisi guncellendi")

    # PREGATE_ARSIV dışındaki özel klasörler – tarama sırasında atlanır,
    # "Zorla Yeniden Oku" ile ayrıca işlenir.
    ATLA_KLASORLER = {"OKUNAMAYANLAR", "FARKLI_FORMAT_DOSYALAR", "ISLEM_RAPORLARI"}

    dosya_listesi = []
    for root, dirs, files in os.walk(ana):
        root_path = Path(root)
        # Bu klasörleri alt klasörlerle birlikte atla
        if any(p in ATLA_KLASORLER for p in root_path.parts):
            dirs[:] = []
            continue

        for file in files:
            yol = root_path / file
            if yol.name.startswith("~$"):
                continue
            dosya_listesi.append(yol)

    # Performans: birden fazla dosya varsa paralel işle (OCR/PDF okuma I/O ağırlıklı)
    if len(dosya_listesi) > 1:
        max_worker = min(4, max(1, os.cpu_count() or 1))
        with ThreadPoolExecutor(max_workers=max_worker) as ex:
            future_map = {ex.submit(dosya_isle, str(yol), ana_klasor): yol for yol in dosya_listesi}
            for fut in as_completed(future_map):
                sonuc = fut.result()
                sonuclar.append(sonuc)
                if log_callback:
                    log_callback(f"{sonuc['İşlem Durumu']}: {sonuc['Eski Dosya Adı']} -> {sonuc.get('Yeni Dosya Adı', '')}")
    else:
        for yol in dosya_listesi:
            sonuc = dosya_isle(str(yol), ana_klasor)
            sonuclar.append(sonuc)
            if log_callback:
                log_callback(f"{sonuc['İşlem Durumu']}: {sonuc['Eski Dosya Adı']} -> {sonuc.get('Yeni Dosya Adı', '')}")

    _eski_t9_klasorlerini_temizle(ana_klasor)

    df = pd.DataFrame(sonuclar)
    rapor_klasor = ana / "ISLEM_RAPORLARI"
    rapor_klasor.mkdir(exist_ok=True)
    rapor_adi = f"ISLEM_RAPORU_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    df.to_excel(rapor_klasor / rapor_adi, index=False)
    df.to_excel(rapor_klasor / "ARSIV_INDEX.xlsx", index=False)
    return df


def arama_yap(ana_klasor, sorgu):
    index_yolu = Path(ana_klasor) / "ISLEM_RAPORLARI" / "ARSIV_INDEX.xlsx"
    if not index_yolu.exists():
        return pd.DataFrame()

    df = pd.read_excel(index_yolu).fillna("")

    # Tam dosya yolu sütunu
    def _tam_yol(row):
        klasor = str(row.get("Yeni Klasör", "")).strip()
        ad     = str(row.get("Yeni Dosya Adı", "")).strip()
        return os.path.join(klasor, ad) if klasor and ad else ""
    df["Tam Yol"] = df.apply(_tam_yol, axis=1)

    if not sorgu:
        return df

    q = temizle_ad(sorgu).replace("_", "")

    def satirda_var(row):
        alan = " ".join(str(x) for x in row.values).upper()
        alan_temiz = temizle_ad(alan).replace("_", "")
        return q in alan_temiz

    return df[df.apply(satirda_var, axis=1)]


def ozet_sayilar(ana_klasor):
    ana   = Path(ana_klasor)
    arsiv = ana / "PREGATE_ARSIV"
    sonuc = {}

    # Birden fazla "tur" aynı fiziksel klasöre (ör. T9_GECICI/T9_MUAYENE/T9_ANA → "T9")
    # işaret edebilir; toplamda mükerrer sayım yapmamak için klasör bazında say.
    klasor_sayilari = {}
    for tur, klasor in BELGE_KLASORLERI.items():
        if klasor not in klasor_sayilari:
            p = arsiv / klasor
            klasor_sayilari[klasor] = len([x for x in p.glob("*.*") if x.is_file()]) if p.exists() else 0
        sonuc[tur] = klasor_sayilari[klasor]

    # PREGATE_ARSIV dışındaki klasörler
    for d in ("OKUNAMAYANLAR", "FARKLI_FORMAT_DOSYALAR"):
        p = ana / d
        sonuc[d] = len([x for x in p.glob("*.*") if x.is_file()]) if p.exists() else 0

    sonuc["TOPLAM"] = sum(klasor_sayilari.values()) + sonuc["OKUNAMAYANLAR"] + sonuc["FARKLI_FORMAT_DOSYALAR"]
    return sonuc


def zorla_yeniden_oku(ana_klasor, log_callback=None):
    """OKUNAMAYANLAR ve FARKLI_FORMAT_DOSYALAR klasörlerindeki dosyaları
    yeniden okumaya çalışır. Okunanlar doğru PREGATE_ARSIV alt klasörüne taşınır.
    Okunamayanlar yerinde kalır. Sonuçları ISLEM_RAPORLARI'na kaydeder."""
    ana   = Path(ana_klasor)
    arsiv = klasorleri_hazirla(ana_klasor)
    sonuclar = []

    # Eski yapıdan kalan iç klasörleri de dahil et
    _eski_konumlari_tasi(ana_klasor)

    hedef_klasorler = [
        ana / "OKUNAMAYANLAR",
        ana / "FARKLI_FORMAT_DOSYALAR",
    ]

    for hedef_dir in hedef_klasorler:
        if not hedef_dir.exists():
            continue
        for dosya in hedef_dir.iterdir():
            if not dosya.is_file() or dosya.name.startswith("~$"):
                continue
            uzanti = dosya.suffix.lower()
            sonuc = {
                "Eski Dosya Adı": dosya.name,
                "Yeni Dosya Adı": dosya.name,
                "Belge Türü": "",
                "Plaka": "", "Tank No": "", "Şasi No": "", "Yabancı Plaka": "",
                "Geçerlilik Tarihi": "", "Kapasite": "", "Tank Kodu": "", "Sürücü": "",
                "Bulunduğu Klasör": str(dosya.parent),
                "Yeni Klasör": str(dosya.parent),
                "İşlem Durumu": "",
                "Hata": "",
                "İşlem Tarihi": datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
            }

            # PRGT etiketliyse → doğru klasöre taşı
            if _is_prgt_dosya(dosya.name):
                tur = belge_turu_dosya_adindan(dosya.name) or "DIGER_BELGELER"
                hk = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
                yol = dosya_tasi(dosya, hk, dosya.name)
                sonuc.update({"Belge Türü": tur, "Yeni Dosya Adı": yol.name,
                              "Yeni Klasör": str(yol.parent), "İşlem Durumu": "PRGT_TASINDI"})
                sonuclar.append(sonuc)
                if log_callback:
                    log_callback(f"PRGT_TASINDI: {dosya.name} → {tur}")
                continue

            # PDF / Resim okumayı dene
            metin = ""
            if desteklenen_uzanti(uzanti):
                try:
                    metin = belge_metni_oku(str(dosya), uzanti)
                except Exception as e:
                    sonuc["Hata"] = str(e)
            else:
                sonuc["İşlem Durumu"] = "DESTEKLENMEYEN_FORMAT"
                sonuclar.append(sonuc)
                if log_callback:
                    log_callback(f"DESTEKLENMEYEN_FORMAT: {dosya.name}")
                continue

            if metin:
                tur = belge_turu_bul(metin, dosya.name)
                bilgi = bilgi_cek(metin, dosya.name)
                yeni_ad = yeni_dosya_adi_olustur(tur, bilgi, dosya.name, uzanti)
                hk = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
                yol = dosya_tasi(dosya, hk, yeni_ad)
                sonuc.update({
                    "Yeni Dosya Adı": yol.name,
                    "Belge Türü": tur,
                    "Plaka": bilgi.get("plaka", ""),
                    "Tank No": bilgi.get("tank_no", ""),
                    "Şasi No": bilgi.get("sasi", ""),
                    "Yabancı Plaka": bilgi.get("yabanci_plaka", ""),
                    "Geçerlilik Tarihi": bilgi.get("tarih", ""),
                    "Kapasite": bilgi.get("kapasite", ""),
                    "Tank Kodu": bilgi.get("tank_kodu", ""),
                    "Sürücü": bilgi.get("surucu", ""),
                    "Yeni Klasör": str(yol.parent),
                    "İşlem Durumu": "ZORLA_OKUNDU_TASINDI",
                })
                if log_callback:
                    log_callback(f"ZORLA_OKUNDU_TASINDI: {dosya.name} → {tur}")
            else:
                sonuc["İşlem Durumu"] = "HALA_OKUNAMADI"
                if log_callback:
                    log_callback(f"HALA_OKUNAMADI: {dosya.name}")

            sonuclar.append(sonuc)

    if sonuclar:
        df = pd.DataFrame(sonuclar)
        rapor_klasor = ana / "ISLEM_RAPORLARI"
        rapor_klasor.mkdir(exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        df.to_excel(rapor_klasor / f"ZORLA_OKUMA_RAPORU_{ts}.xlsx", index=False)
        return df
    return pd.DataFrame()


def _yeniden_oku_ve_isimlendir(kaynak, arsiv, eski_ad, uzanti):
    """İçeriği baştan okuyup sınıflandırarak yeni 'S' kodlu adla doğru
    (birleşik T9 dahil) klasöre taşır. Okunamazsa OKUNAMAYANLAR'a, desteklenmeyen
    formatsa FARKLI_FORMAT_DOSYALAR'a taşır. Sonuc sözlüğünü döndürür."""
    sonuc_guncelleme = {}

    if not desteklenen_uzanti(uzanti):
        yeni_yol = dosya_tasi(kaynak, arsiv.parent / "FARKLI_FORMAT_DOSYALAR")
        return {
            "Belge Türü": "FARKLI_FORMAT_DOSYALAR", "Yeni Dosya Adı": yeni_yol.name,
            "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_FARKLI_FORMAT",
        }

    try:
        metin = belge_metni_oku(str(kaynak), uzanti)
    except Exception as e:
        metin = ""
        sonuc_guncelleme["Hata"] = f"OKUMA_HATASI: {e}"

    tur = belge_turu_bul(metin, eski_ad)

    if not metin and tur == "DIGER_BELGELER":
        yeni_yol = dosya_tasi(kaynak, arsiv.parent / "OKUNAMAYANLAR")
        sonuc_guncelleme.update({
            "Belge Türü": "OKUNAMAYANLAR", "Yeni Dosya Adı": yeni_yol.name,
            "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_OKUNAMADI",
        })
        return sonuc_guncelleme

    bilgi = bilgi_cek(metin, eski_ad)
    yeni_ad = yeni_dosya_adi_olustur(tur, bilgi, eski_ad, uzanti)
    yeni_yol = dosya_tasi(kaynak, arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER"), yeni_ad)

    sonuc_guncelleme.update({
        "Yeni Dosya Adı": yeni_yol.name, "Belge Türü": tur,
        "Plaka": bilgi.get("plaka", ""), "Tank No": bilgi.get("tank_no", ""),
        "Şasi No": bilgi.get("sasi", ""), "Yabancı Plaka": bilgi.get("yabanci_plaka", ""),
        "Geçerlilik Tarihi": bilgi.get("tarih", ""), "Kapasite": bilgi.get("kapasite", ""),
        "Tank Kodu": bilgi.get("tank_kodu", ""), "Sürücü": bilgi.get("surucu", ""),
        "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_OKUNDU",
    })
    return sonuc_guncelleme


def _dosya_yeniden_isle(dosya_yolu, ana_klasor, arsiv):
    """'Tümünü Yeniden Tara' için tek dosya işleme.

    - OKUNAMAYANLAR / FARKLI_FORMAT_DOSYALAR içindeki dosyalar: içerik yeniden
      okunmaya çalışılır (zorla yeniden oku ile aynı mantık), başarılı olursa
      'S' koduyla doğru klasöre taşınır.
    - PREGATE_ARSIV içinde, kodu 'E' / eski 'PRGT' olan dosyalar: içerik tekrar
      okunmaz; sadece kodu 'E' yapılır ve doğru (birleşik T9 dahil) klasöre taşınır.
    - PREGATE_ARSIV içinde, kodu 'S' / eski 'PRGTOK' olan dosyalar: içerik tekrar
      okunmaz (zaten doğru sınıflandırılmış); sadece kodu 'S' yapılır ve
      birleşik T9 dahil doğru klasöre taşınır.
    - Hiç kodu olmayan dosyalar: içerik baştan okunur ve sınıflandırılır."""
    kaynak = Path(dosya_yolu)
    ana = Path(ana_klasor)
    eski_ad, uzanti = kaynak.name, kaynak.suffix.lower()

    sonuc = {
        "Eski Dosya Adı": eski_ad, "Yeni Dosya Adı": "", "Belge Türü": "", "Plaka": "", "Tank No": "",
        "Şasi No": "", "Yabancı Plaka": "", "Geçerlilik Tarihi": "", "Kapasite": "", "Tank Kodu": "",
        "Sürücü": "", "Bulunduğu Klasör": str(kaynak.parent), "Yeni Klasör": "", "İşlem Durumu": "",
        "Hata": "", "İşlem Tarihi": datetime.now().strftime("%d-%m-%Y %H:%M:%S"),
    }

    try:
        if kaynak.name.startswith("~$") or not kaynak.exists():
            sonuc["İşlem Durumu"] = "GECICI_DOSYA_ATLANDI"
            return sonuc

        okunamayan_klasorler = (
            (ana / "OKUNAMAYANLAR").resolve(),
            (ana / "FARKLI_FORMAT_DOSYALAR").resolve(),
        )

        # ── OKUNAMAYANLAR / FARKLI_FORMAT_DOSYALAR: yeniden okumayı dene ─────────
        if kaynak.parent.resolve() in okunamayan_klasorler:
            if _is_prgt_dosya(kaynak.name):
                tur = belge_turu_dosya_adindan(kaynak.name) or "DIGER_BELGELER"
                hedef_klasor = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
                yeni_ad = kaynak.name if _son_token(kaynak.name) == ELLE_DUZELT_KODU else _kod_degistir(kaynak.name, ELLE_DUZELT_KODU)
                yeni_yol = dosya_tasi(kaynak, hedef_klasor, yeni_ad)
                sonuc.update({
                    "Yeni Dosya Adı": yeni_yol.name, "Belge Türü": tur,
                    "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_ELLEDUZ",
                })
                return sonuc

            sonuc.update(_yeniden_oku_ve_isimlendir(kaynak, arsiv, eski_ad, uzanti))
            if not sonuc.get("İşlem Durumu"):
                sonuc["İşlem Durumu"] = "TAM_TARAMA_HALA_OKUNAMADI"
            return sonuc

        # ── PREGATE_ARSIV içinde, elle düzeltilmiş (E / eski PRGT) ───────────────
        if _is_prgt_dosya(kaynak.name):
            tur = belge_turu_dosya_adindan(kaynak.name) or "DIGER_BELGELER"
            hedef_klasor = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
            yeni_ad = kaynak.name if _son_token(kaynak.name) == ELLE_DUZELT_KODU else _kod_degistir(kaynak.name, ELLE_DUZELT_KODU)
            yeni_yol = dosya_tasi(kaynak, hedef_klasor, yeni_ad)
            sonuc.update({
                "Yeni Dosya Adı": yeni_yol.name, "Belge Türü": tur,
                "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_ELLEDUZ",
            })
            return sonuc

        # ── PREGATE_ARSIV içinde, daha önce işlenmiş (S / eski PRGTOK) ───────────
        if _islenmis_mi(kaynak.name):
            tur = belge_turu_dosya_adindan(kaynak.name) or "DIGER_BELGELER"
            hedef_klasor = arsiv / BELGE_KLASORLERI.get(tur, "DIGER_BELGELER")
            yeni_ad = kaynak.name if _son_token(kaynak.name) == ISLEM_KODU else _kod_degistir(kaynak.name, ISLEM_KODU)
            yeni_yol = dosya_tasi(kaynak, hedef_klasor, yeni_ad)
            sonuc.update({
                "Yeni Dosya Adı": yeni_yol.name, "Belge Türü": tur,
                "Yeni Klasör": str(yeni_yol.parent), "İşlem Durumu": "TAM_TARAMA_KOD_GUNCELLENDI",
            })
            return sonuc

        # ── Kodu olmayan dosya: içerik baştan okunur ve sınıflandırılır ──────────
        sonuc.update(_yeniden_oku_ve_isimlendir(kaynak, arsiv, eski_ad, uzanti))
        return sonuc

    except Exception as e:
        sonuc["İşlem Durumu"] = "HATA"
        sonuc["Hata"] = str(e)
        return sonuc


def tum_dosyalari_yeniden_tara(ana_klasor, log_callback=None):
    """TÜM arşivdeki dosyaları (PREGATE_ARSIV + OKUNAMAYANLAR + FARKLI_FORMAT_DOSYALAR)
    içeriklerini yeniden okuyarak baştan sona tekrar tarar.

    - Eski 'PRGTOK' / 'PRGT' kodlu dosyalar yeni 'S' / 'E' koduna çevrilir.
    - T9_GECICI / T9_MUAYENE / T9_ANA / T9_MUAYENE_SERTIFIKASI klasörlerindeki
      dosyalar tek bir 'T9' klasöründe birleştirilir.
    - Bu işlem normal taramadan çok daha uzun sürebilir; bu yüzden ayrı bir
      buton ile elle tetiklenir."""
    ana   = Path(ana_klasor)
    arsiv = klasorleri_hazirla(ana_klasor)
    sonuclar = []

    _eski_konumlari_tasi(ana_klasor)
    if log_callback:
        log_callback("Eski klasor yapisi guncellendi")

    taranacak_dizinler = [arsiv, ana / "OKUNAMAYANLAR", ana / "FARKLI_FORMAT_DOSYALAR"]

    dosya_listesi = []
    for taban in taranacak_dizinler:
        if not taban.exists():
            continue
        for root, dirs, files in os.walk(taban):
            root_path = Path(root)
            if "ISLEM_RAPORLARI" in root_path.parts:
                continue
            for file in files:
                yol = root_path / file
                if yol.name.startswith("~$"):
                    continue
                dosya_listesi.append(yol)

    if len(dosya_listesi) > 1:
        max_worker = min(4, max(1, os.cpu_count() or 1))
        with ThreadPoolExecutor(max_workers=max_worker) as ex:
            future_map = {ex.submit(_dosya_yeniden_isle, yol, ana_klasor, arsiv): yol for yol in dosya_listesi}
            for fut in as_completed(future_map):
                sonuc = fut.result()
                sonuclar.append(sonuc)
                if log_callback:
                    log_callback(f"{sonuc['İşlem Durumu']}: {sonuc['Eski Dosya Adı']} -> {sonuc.get('Yeni Dosya Adı', '')}")
    else:
        for yol in dosya_listesi:
            sonuc = _dosya_yeniden_isle(yol, ana_klasor, arsiv)
            sonuclar.append(sonuc)
            if log_callback:
                log_callback(f"{sonuc['İşlem Durumu']}: {sonuc['Eski Dosya Adı']} -> {sonuc.get('Yeni Dosya Adı', '')}")

    _eski_t9_klasorlerini_temizle(ana_klasor)

    df = pd.DataFrame(sonuclar)
    rapor_klasor = ana / "ISLEM_RAPORLARI"
    rapor_klasor.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    df.to_excel(rapor_klasor / f"TAM_YENIDEN_TARAMA_RAPORU_{ts}.xlsx", index=False)
    if not df.empty:
        df.to_excel(rapor_klasor / "ARSIV_INDEX.xlsx", index=False)
    return df
